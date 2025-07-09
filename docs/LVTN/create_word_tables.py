#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Script tạo file Word với bảng từ markdown
Yêu cầu: pip install python-docx
"""

import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

def parse_markdown_tables():
    """Parse markdown file và trích xuất thông tin bảng"""
    
    with open('description.md', 'r', encoding='utf-8') as f:
        content = f.read()
    
    tables = []
    lines = content.split('\n')
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        # Tìm tiêu đề bảng
        if line.startswith('## Bảng '):
            table_name = line.replace('## Bảng ', '').strip()
            
            # Tìm mô tả
            i += 1
            description = ""
            while i < len(lines) and not lines[i].strip().startswith('|'):
                if lines[i].strip().startswith('**Mô tả:**'):
                    description = lines[i].strip().replace('**Mô tả:**', '').strip()
                i += 1
            
            # Parse bảng markdown
            columns = []
            if i < len(lines) and '| Thuộc tính' in lines[i]:
                i += 1  # Skip header
                if i < len(lines) and ('|---' in lines[i] or '| :---' in lines[i]):
                    i += 1  # Skip separator
                
                # Đọc dữ liệu
                while i < len(lines):
                    data_line = lines[i].strip()
                    if not data_line or not data_line.startswith('|'):
                        break
                    
                    parts = [p.strip() for p in data_line.split('|')[1:-1]]
                    if len(parts) >= 6:
                        columns.append({
                            'attribute': parts[0],
                            'type': parts[1], 
                            'key': parts[2],
                            'unique': parts[3],
                            'mandatory': parts[4],
                            'description': parts[5]
                        })
                    i += 1
            
            tables.append({
                'name': table_name,
                'description': description,
                'columns': columns
            })
        else:
            i += 1
    
    return tables

def set_cell_border(cell, **kwargs):
    """Thêm border cho cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Tạo borders element
    tcBorders = OxmlElement('w:tcBorders')
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_data['val'])
            element.set(qn('w:sz'), str(edge_data['sz']))
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), edge_data['color'])
            tcBorders.append(element)
    
    tcPr.append(tcBorders)

def create_word_document(tables):
    """Tạo document Word với bảng"""
    
    doc = Document()
    
    # Tiêu đề chính
    title = doc.add_heading('Mô tả Database - ThanhHuyStore E-commerce Platform', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Tạo bảng tổng quan
    doc.add_heading('Tổng quan các bảng', level=1)
    
    overview_table = doc.add_table(rows=1, cols=4)
    overview_table.style = 'Table Grid'
    overview_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header cho bảng tổng quan
    overview_headers = ['STT', 'Tên Bảng', 'Mô tả', 'Số cột']
    for i, header in enumerate(overview_headers):
        cell = overview_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thêm dữ liệu tổng quan
    for idx, table in enumerate(tables, 1):
        row = overview_table.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = table['name']
        row.cells[2].text = table['description']
        row.cells[3].text = str(len(table['columns']))
        
        # Center align STT và số cột
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Set column widths cho bảng tổng quan
    for row in overview_table.rows:
        row.cells[0].width = Inches(0.5)  # STT
        row.cells[1].width = Inches(2.0)  # Tên bảng
        row.cells[2].width = Inches(4.0)  # Mô tả
        row.cells[3].width = Inches(0.8)  # Số cột
    
    doc.add_page_break()
    
    # Tạo bảng chi tiết cho từng bảng
    for table in tables:
        # Tiêu đề bảng
        heading = doc.add_heading(f'Bảng {table["name"]}', level=1)
        
        # Mô tả bảng
        desc_para = doc.add_paragraph()
        desc_para.add_run('Mô tả: ').bold = True
        desc_para.add_run(table['description'])
        
        # Tạo bảng chi tiết
        detail_table = doc.add_table(rows=1, cols=6)
        detail_table.style = 'Table Grid'
        detail_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header
        headers = ['Thuộc tính', 'Kiểu dữ liệu', 'Khóa', 'Duy nhất', 'Bắt buộc', 'Diễn giải']
        for i, header in enumerate(headers):
            cell = detail_table.rows[0].cells[i]
            cell.text = header
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Thêm dữ liệu
        for column in table['columns']:
            row = detail_table.add_row()
            row.cells[0].text = column['attribute']
            row.cells[1].text = column['type']
            row.cells[2].text = column['key']
            row.cells[3].text = column['unique']
            row.cells[4].text = column['mandatory']
            row.cells[5].text = column['description']
            
            # Center align các cột ngắn
            for i in [2, 3, 4]:  # Khóa, Duy nhất, Bắt buộc
                row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Set column widths
        for row in detail_table.rows:
            row.cells[0].width = Inches(2.0)   # Thuộc tính
            row.cells[1].width = Inches(1.2)   # Kiểu dữ liệu
            row.cells[2].width = Inches(0.6)   # Khóa
            row.cells[3].width = Inches(0.8)   # Duy nhất
            row.cells[4].width = Inches(0.8)   # Bắt buộc
            row.cells[5].width = Inches(3.0)   # Diễn giải
        
        doc.add_paragraph()  # Thêm khoảng trống
    
    return doc

def create_simple_text_format(tables):
    """Tạo file text đơn giản để copy vào Word"""
    
    content = []
    content.append("MÔ TẢ DATABASE - THANHHUYSTORE E-COMMERCE PLATFORM")
    content.append("=" * 60)
    content.append("")
    
    # Tổng quan
    content.append("TỔNG QUAN CÁC BẢNG")
    content.append("-" * 30)
    content.append("")
    
    for idx, table in enumerate(tables, 1):
        content.append(f"{idx:2d}. {table['name']:<20} - {len(table['columns']):2d} cột")
        content.append(f"    {table['description']}")
        content.append("")
    
    content.append("\n" + "=" * 60 + "\n")
    
    # Chi tiết từng bảng
    for table in tables:
        content.append(f"BẢNG: {table['name'].upper()}")
        content.append("-" * 40)
        content.append(f"Mô tả: {table['description']}")
        content.append("")
        
        # Header
        content.append(f"{'Thuộc tính':<25} {'Kiểu':<12} {'K':<3} {'U':<3} {'M':<3} {'Diễn giải'}")
        content.append("-" * 80)
        
        # Data
        for column in table['columns']:
            content.append(f"{column['attribute']:<25} {column['type']:<12} {column['key']:<3} {column['unique']:<3} {column['mandatory']:<3} {column['description']}")
        
        content.append("\n" + "=" * 60 + "\n")
    
    return "\n".join(content)

def main():
    """Hàm chính"""
    print("Đang parse file markdown...")
    tables = parse_markdown_tables()
    print(f"Đã tìm thấy {len(tables)} bảng")
    
    try:
        # Tạo file Word
        print("Đang tạo file Word...")
        doc = create_word_document(tables)
        word_filename = "csv_output/database_description.docx"
        doc.save(word_filename)
        print(f"Đã tạo file Word: {word_filename}")
    except ImportError:
        print("Không thể tạo file Word (cần cài đặt python-docx)")
        print("Chạy: pip install python-docx")
    
    # Tạo file text đơn giản
    print("Đang tạo file text format...")
    text_content = create_simple_text_format(tables)
    text_filename = "csv_output/database_simple_format.txt"
    with open(text_filename, 'w', encoding='utf-8') as f:
        f.write(text_content)
    print(f"Đã tạo file text: {text_filename}")
    
    print("\nHoàn thành! Bạn có thể:")
    print("1. Mở file .docx trực tiếp bằng Word")
    print("2. Copy nội dung từ file .txt và paste vào Word")
    print("3. Sử dụng các file CSV để import vào Excel")

if __name__ == "__main__":
    main()
